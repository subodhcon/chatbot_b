import os
import logging
from typing import Dict, Any, List
from docx import Document

logger = logging.getLogger("app.services.docx_extraction")


class DOCXExtractionService:
    """
    Service for extracting structured text, tables, and metadata from DOCX documents using python-docx.
    """

    def extract_docx_content(self, file_path: str) -> Dict[str, Any]:
        """
        Extracts metadata and text (paragraphs + tables) from a DOCX file.
        Returns a structured dictionary containing metadata, total text, and stats.
        Raises ValueError if the DOCX is corrupted or unreadable.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found at path: {file_path}")

        try:
            doc = Document(file_path)

            # Extract document metadata
            props = doc.core_properties
            metadata = {
                "title": props.title or "",
                "author": props.author or "",
                "subject": props.subject or "",
                "keywords": props.keywords or "",
                "category": props.category or "",
                "comments": props.comments or "",
            }

            text_parts: List[str] = []

            # 1. Extract text from paragraphs
            for paragraph in doc.paragraphs:
                p_text = paragraph.text.strip()
                if p_text:
                    text_parts.append(p_text)

            # 2. Extract text from tables (row by row, cells separated by pipe)
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    # Filter out empty cells, keep meaningful table row entries
                    meaningful_cells = [c for c in row_cells if c]
                    if meaningful_cells:
                        text_parts.append(" | ".join(meaningful_cells))

            total_text = "\n\n".join(text_parts)

            return {
                "metadata": metadata,
                "total_text": total_text,
                "char_count": len(total_text),
                "word_count": len(total_text.split()),
            }

        except Exception as e:
            logger.error(f"Failed to extract text from DOCX '{file_path}': {e}", exc_info=True)
            raise ValueError(f"Failed to parse DOCX file. It may be corrupted or invalid. Error: {str(e)}") from e


# Module-level singleton
docx_extraction_service = DOCXExtractionService()
